"""
Reusable module + CLI for creating Word (.docx) documents with python-docx.

Module API (import and call from Python):
    doc = new_doc(output_path, title=..., font_name=..., font_size_pt=..., margin_inches=...)
    add_title(doc, text)
    add_subtitle(doc, text)
    add_section_header(doc, text, color=COLOR_WALMART_BLUE)
    add_subheader(doc, text, color=COLOR_WALMART_BLUE)
    add_divider(doc)
    add_heading(doc, text, level=1)
    add_paragraph(doc, text, bold=False, italic=False)
    add_bullet(doc, text, level=0)
    add_table(doc, headers, rows)
    add_page_break(doc)
    path = save_doc(doc, output_path)

CLI (for quick testing):
    python3 scripts/make_doc.py demo --output /tmp/test.docx
"""

from __future__ import annotations

import argparse
import logging
import os
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

DEFAULT_FONT = "Everyday Sans"
DEFAULT_FONT_SIZE = 10.0
DEFAULT_MARGIN = 0.5

# ---------------------------------------------------------------------------
# Color constants
# ---------------------------------------------------------------------------

COLOR_WALMART_BLUE = "003594"   # primary headers
COLOR_HEADING_1 = "365F91"      # Heading 1
COLOR_HEADING_2 = "4F81BD"      # Heading 2/3
COLOR_GRAY = "646464"           # subtitles
COLOR_GREEN = "146E37"          # accent/callout
COLOR_DIVIDER = "C8C8C8"        # divider line


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _rgb(hex_color: str) -> RGBColor:
    """Convert a 6-char hex color string (no leading #) to RGBColor."""
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def _apply_font(
    run,
    name: str,
    size_pt: float,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Apply font name, size, bold, and italic to a run.

    Also patches the low-level XML rFonts element so Word respects the font
    even when theme fonts would otherwise override it (the same technique used
    in the work-meetings scripts).
    """
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic

    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), name)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


# ---------------------------------------------------------------------------
# Public module API
# ---------------------------------------------------------------------------

def new_doc(
    output_path: str,
    title: Optional[str] = None,
    font_name: str = DEFAULT_FONT,
    font_size_pt: float = DEFAULT_FONT_SIZE,
    margin_inches: float = DEFAULT_MARGIN,
) -> Document:
    """Create a new Document with default formatting applied.

    Configures Normal style and Heading 1-3 styles with branded colors and
    spacing so add_heading() picks them up automatically.

    Args:
        output_path: Where the file will be saved (used for logging only here;
                     actual save happens in save_doc).
        title: Optional document title added at the top via add_title().
        font_name: Body font name. Default: "Everyday Sans".
        font_size_pt: Body font size in points. Default: 10.0.
        margin_inches: All four margins in inches. Default: 0.5.

    Returns:
        A python-docx Document object ready to be populated.
    """
    logger.debug("Creating new document: output_path=%s", output_path)
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)

    # Normal style baseline
    doc.styles["Normal"].font.name = font_name
    doc.styles["Normal"].font.size = Pt(font_size_pt)

    # Heading 1: 14pt, bold, #365F91, space_before=24pt, space_after=0pt
    h1 = doc.styles["Heading 1"]
    h1.font.name = font_name
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = _rgb(COLOR_HEADING_1)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(0)

    # Heading 2: 13pt, bold, #4F81BD, space_before=10pt, space_after=0pt
    h2 = doc.styles["Heading 2"]
    h2.font.name = font_name
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = _rgb(COLOR_HEADING_2)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(0)

    # Heading 3: 10pt, bold, #4F81BD, space_before=10pt, space_after=0pt
    h3 = doc.styles["Heading 3"]
    h3.font.name = font_name
    h3.font.size = Pt(10)
    h3.font.bold = True
    h3.font.color.rgb = _rgb(COLOR_HEADING_2)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(0)

    # Store formatting on the doc object so the other helpers can read it
    doc._yuki_font_name = font_name  # type: ignore[attr-defined]
    doc._yuki_font_size = font_size_pt  # type: ignore[attr-defined]

    if title:
        add_title(doc, title)

    return doc


def add_title(doc: Document, text: str) -> None:
    """Add a document title: 15pt bold #003594. space_after=2pt.

    Args:
        doc: Document returned by new_doc().
        text: Title text.
    """
    font_name: str = getattr(doc, "_yuki_font_name", DEFAULT_FONT)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _apply_font(run, font_name, 15.0, bold=True)
    run.font.color.rgb = _rgb(COLOR_WALMART_BLUE)
    logger.debug("add_title: %s", text)


def add_subtitle(doc: Document, text: str) -> None:
    """Add a subtitle line: 10pt #646464. space_after=8pt.

    Args:
        doc: Document returned by new_doc().
        text: Subtitle text.
    """
    font_name: str = getattr(doc, "_yuki_font_name", DEFAULT_FONT)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    _apply_font(run, font_name, 10.0)
    run.font.color.rgb = _rgb(COLOR_GRAY)
    logger.debug("add_subtitle: %s", text)


def add_section_header(doc: Document, text: str, color: str = COLOR_WALMART_BLUE) -> None:
    """Add an ALL-CAPS bold section header.

    12pt, Walmart blue (#003594) by default. space_before=10pt, space_after=3pt.
    Text is uppercased automatically.

    Args:
        doc: Document returned by new_doc().
        text: Header text (will be uppercased automatically).
        color: Hex color string (no #). Default: COLOR_WALMART_BLUE.
    """
    font_name: str = getattr(doc, "_yuki_font_name", DEFAULT_FONT)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    _apply_font(run, font_name, 12.0, bold=True)
    run.font.color.rgb = _rgb(color)
    logger.debug("add_section_header: %s", text)


def add_subheader(doc: Document, text: str, color: str = COLOR_WALMART_BLUE) -> None:
    """Add a bold sub-section header.

    Inherits default font size, color #003594. space_before=5pt, space_after=2pt.

    Args:
        doc: Document returned by new_doc().
        text: Subheader text.
        color: Hex color string (no #). Default: COLOR_WALMART_BLUE.
    """
    font_name: str = getattr(doc, "_yuki_font_name", DEFAULT_FONT)
    font_size: float = getattr(doc, "_yuki_font_size", DEFAULT_FONT_SIZE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _apply_font(run, font_name, font_size, bold=True)
    run.font.color.rgb = _rgb(color)
    logger.debug("add_subheader: %s", text)


def add_divider(doc: Document) -> None:
    """Add a decorative horizontal divider.

    A row of em dashes (──────────) in 7pt #C8C8C8. space_before=3pt, space_after=3pt.

    Args:
        doc: Document returned by new_doc().
    """
    font_name: str = getattr(doc, "_yuki_font_name", DEFAULT_FONT)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("\u2500" * 40)
    _apply_font(run, font_name, 7.0)
    run.font.color.rgb = _rgb(COLOR_DIVIDER)
    logger.debug("add_divider")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading paragraph using the Word Heading style.

    Style colors and spacing are configured by new_doc(). Run-level font
    is patched to ensure the correct font name is respected by Word.

    Level 1 → 14pt #365F91, space_before=24pt
    Level 2 → 13pt #4F81BD, space_before=10pt
    Level 3 → 10pt #4F81BD, space_before=10pt

    Args:
        doc: Document returned by new_doc().
        text: Heading text.
        level: Heading level 1-3.
    """
    font_name: str = getattr(doc, "_yuki_font_name", DEFAULT_FONT)

    size_map = {1: 14.0, 2: 13.0, 3: 10.0}
    size_pt = size_map.get(level, 14.0)
    color_map = {1: COLOR_HEADING_1, 2: COLOR_HEADING_2, 3: COLOR_HEADING_2}
    color = color_map.get(level, COLOR_HEADING_1)

    p = doc.add_paragraph(style="Heading %d" % level)
    run = p.add_run(text)
    _apply_font(run, font_name, size_pt, bold=True)
    run.font.color.rgb = _rgb(color)
    logger.debug("add_heading level=%d: %s", level, text)


def add_paragraph(
    doc: Document,
    text: str,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Add a body paragraph with default font applied.

    Args:
        doc: Document returned by new_doc().
        text: Paragraph text.
        bold: Whether to bold the text.
        italic: Whether to italicize the text.
    """
    font_name: str = getattr(doc, "_yuki_font_name", DEFAULT_FONT)
    font_size: float = getattr(doc, "_yuki_font_size", DEFAULT_FONT_SIZE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _apply_font(run, font_name, font_size, bold=bold, italic=italic)
    logger.debug("add_paragraph bold=%s italic=%s: %s", bold, italic, text[:60])


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    """Add a bullet-list item.

    Args:
        doc: Document returned by new_doc().
        text: Bullet text.
        level: 0 = top-level bullet, 1 = indented sub-bullet.
    """
    font_name: str = getattr(doc, "_yuki_font_name", DEFAULT_FONT)
    font_size: float = getattr(doc, "_yuki_font_size", DEFAULT_FONT_SIZE)

    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Inches(0.2 + level * 0.2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    _apply_font(run, font_name, font_size)
    logger.debug("add_bullet level=%d: %s", level, text[:60])


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    """Add a simple table with a header row.

    The header row is bold. All cells use the default document font.

    Args:
        doc: Document returned by new_doc().
        headers: Column header strings.
        rows: Data rows — each row is a list of cell strings.
    """
    font_name: str = getattr(doc, "_yuki_font_name", DEFAULT_FONT)
    font_size: float = getattr(doc, "_yuki_font_size", DEFAULT_FONT_SIZE)

    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # Header row
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(header_text)
        _apply_font(run, font_name, font_size, bold=True)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = ""
            run = cell.paragraphs[0].add_run(cell_text)
            _apply_font(run, font_name, font_size)

    logger.debug("add_table %d cols x %d data rows", col_count, len(rows))


def add_page_break(doc: Document) -> None:
    """Insert a page break.

    Args:
        doc: Document returned by new_doc().
    """
    doc.add_page_break()
    logger.debug("add_page_break")


def save_doc(doc: Document, output_path: str) -> str:
    """Save the document and return the absolute path.

    Args:
        doc: Document returned by new_doc().
        output_path: File path to save to (created or overwritten).

    Returns:
        Absolute path of the saved file.
    """
    abs_path = os.path.abspath(output_path)
    doc.save(abs_path)
    logger.info("Saved document: %s", abs_path)
    return abs_path


# ---------------------------------------------------------------------------
# CLI demo command
# ---------------------------------------------------------------------------

def _run_demo(output_path: str) -> None:
    """Create a sample document exercising all features."""
    doc = new_doc(output_path)

    add_title(doc, "Word Document Skill — Demo")
    add_subtitle(doc, "Generated by make_doc.py · Branded formatting test")
    add_divider(doc)

    add_section_header(doc, "Typography & Headings")
    add_heading(doc, "Section 1: Heading Level 1", level=1)
    add_paragraph(doc, "Normal body text at 10 pt Everyday Sans.")
    add_paragraph(doc, "Bold body text.", bold=True)
    add_paragraph(doc, "Italic body text.", italic=True)

    add_heading(doc, "Subsection 1.1: Heading Level 2", level=2)
    add_subheader(doc, "Subheader using add_subheader()")
    add_paragraph(doc, "Content under a subheader.")

    add_heading(doc, "1.1.1 Heading Level 3", level=3)
    add_paragraph(doc, "Content under a level-3 heading.")

    add_divider(doc)

    add_section_header(doc, "Bullets")
    add_bullet(doc, "Top-level bullet item A")
    add_bullet(doc, "Indented sub-bullet under A", level=1)
    add_bullet(doc, "Top-level bullet item B")
    add_bullet(doc, "Another sub-bullet", level=1)

    add_section_header(doc, "Tables")
    add_paragraph(doc, "A simple three-column table:")
    add_table(
        doc,
        headers=["Name", "Role", "Status"],
        rows=[
            ["Yuki", "AI Assistant", "Active"],
            ["Claude", "Underlying Model", "Available"],
            ["python-docx", "Library", "1.2.0"],
        ],
    )

    add_page_break(doc)

    add_title(doc, "Page 2 — Color Palette")
    add_subtitle(doc, "All color constants demonstrated below")
    add_divider(doc)

    add_section_header(doc, "Walmart Blue (default, #003594)")
    add_section_header(doc, "Green accent (#146E37)", color=COLOR_GREEN)
    add_subheader(doc, "Subheader in Walmart Blue (default)")
    add_paragraph(doc, "This content appears on the second page.")

    path = save_doc(doc, output_path)
    logger.info("Saved: %s", path)


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(name)s %(message)s")

    parser = argparse.ArgumentParser(description="make_doc.py — Word document helper")
    subparsers = parser.add_subparsers(dest="command", required=True)

    demo_parser = subparsers.add_parser("demo", help="Generate a sample .docx exercising all features")
    demo_parser.add_argument("--output", required=True, help="Output path, e.g. /tmp/test.docx")
    demo_parser.add_argument("--font", default=DEFAULT_FONT, help="Font name (default: Everyday Sans)")
    demo_parser.add_argument("--font-size", type=float, default=DEFAULT_FONT_SIZE, help="Font size in pt (default: 10.0)")
    demo_parser.add_argument("--margin", type=float, default=DEFAULT_MARGIN, help="Margin in inches (default: 0.5)")

    args = parser.parse_args()

    if args.command == "demo":
        _run_demo(args.output)


if __name__ == "__main__":
    main()
